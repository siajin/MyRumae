"""User-facing DOCX summary generator (정리/ folder).

Post-refactor layout (one DOCX per source file/notice/assignment, grouped by
source activity rather than week):

    Desktop/UOS_LMS_AI/<course>/<source_label>/정리/
        <file_stem>.docx            ← per Material (PDF/PPT/DOCX parsed body)
        공지_<notice_title>.docx     ← per Notice (board article body)
        과제_<assign_title>.docx     ← per Assignment (description + due_at)

The "파싱 본문" section embeds the actual extracted text from ParsedContent in
reading order, prefixed by page numbers so the user can cross-reference the
original PDF/slide deck.

AI summary branch is kept dormant — `Summary` table / `repo.upsert_summary`
exist but are not read here.
"""
from __future__ import annotations

import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlalchemy.orm import Session

from ..db.models import (
    Assignment,
    Course,
    CourseSyllabus,
    Material,
    Notice,
    ParsedContent,
    Timetable,
)
from ..downloader import paths

log = logging.getLogger(__name__)

_MAX_FILENAME_LEN = 80


# ---------------------------------------------------------------------------
# Queries

def _materials_for_course(db: Session, course_id: int) -> List[Material]:
    return (
        db.query(Material)
        .filter(Material.course_id == course_id)
        .order_by(Material.uploaded_at.asc())
        .all()
    )


def _notices_for_course(db: Session, course_id: int) -> List[Notice]:
    return (
        db.query(Notice)
        .filter(Notice.course_id == course_id)
        .order_by(Notice.posted_at.asc())
        .all()
    )


def _assignments_for_course(db: Session, course_id: int) -> List[Assignment]:
    return (
        db.query(Assignment)
        .filter(Assignment.course_id == course_id)
        .order_by(Assignment.due_at.asc())
        .all()
    )


def _parsed_for_material(db: Session, material_id: int) -> Optional[ParsedContent]:
    return (
        db.query(ParsedContent)
        .filter(ParsedContent.material_id == material_id)
        .one_or_none()
    )


def _timetable_for_course(db: Session, course_id: int) -> List[Timetable]:
    return (
        db.query(Timetable)
        .filter(Timetable.course_id == course_id)
        .order_by(Timetable.weekday.asc(), Timetable.start_time.asc())
        .all()
    )


def _course_week_topic(db: Session, course_id: int, week: Optional[int]) -> Optional[dict]:
    if week is None:
        return None
    syl = (
        db.query(CourseSyllabus)
        .filter(CourseSyllabus.course_id == course_id)
        .one_or_none()
    )
    if syl is None or not syl.schedule_json:
        return None
    try:
        weeks = json.loads(syl.schedule_json)
    except Exception:
        log.exception("schedule_json unreadable for course %d", course_id)
        return None
    for w in weeks:
        if w.get("week") == week:
            return w
    return None


def _load_blocks(pc: ParsedContent) -> list[dict]:
    if not pc.blocks_json_path:
        return []
    p = Path(pc.blocks_json_path)
    if not p.exists():
        return []
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        log.exception("blocks json unreadable: %s", p)
        return []


# ---------------------------------------------------------------------------
# DOCX building blocks

def _set_korean_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)


def _add_meta(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(value)


def _add_section_heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _add_bullet(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    if bold:
        run.bold = True


_TAG_RE = re.compile(r"<[^>]+>")
_WS_RE = re.compile(r"[ \t\f\v]+")
_BLANK_LINE_RE = re.compile(r"\n\s*\n+")


def _html_to_text(html: Optional[str]) -> str:
    if not html:
        return ""
    # <br> and block-ish closers become newlines so paragraphs survive.
    text = re.sub(r"(?i)<\s*br\s*/?\s*>", "\n", html)
    text = re.sub(r"(?i)</\s*(p|div|li|tr|h[1-6])\s*>", "\n", text)
    text = _TAG_RE.sub("", text)
    text = text.replace("\r", "")
    text = _WS_RE.sub(" ", text)
    text = _BLANK_LINE_RE.sub("\n\n", text).strip()
    return text


def _safe_filename(stem: str, fallback: str = "untitled") -> str:
    clean = paths.sanitize_segment(stem or fallback, fallback=fallback)
    if len(clean) > _MAX_FILENAME_LEN:
        clean = clean[:_MAX_FILENAME_LEN].rstrip(" ._-")
    return clean or fallback


# ---------------------------------------------------------------------------
# Helpers shared by the per-file generators

def _human_size(n: Optional[int]) -> str:
    if not n:
        return ""
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024:
            return f"{n:.0f}{unit}" if unit == "B" else f"{n:.1f}{unit}"
        n /= 1024
    return f"{n:.1f}TB"


_WEEKDAY_KO = ("월", "화", "수", "목", "금", "토", "일")


def _format_slot(slot: Timetable) -> str:
    day = _WEEKDAY_KO[slot.weekday] if 0 <= slot.weekday < 7 else "?"
    times = slot.start_time or ""
    if slot.end_time:
        times = f"{times}–{slot.end_time}"
    loc = f" ({slot.location})" if slot.location else ""
    return f"{day} {times}{loc}".strip()


def _course_header(doc: Document, course: Course, source_label: Optional[str]) -> None:
    _add_meta(doc, "과목", course.course_name)
    if course.course_code:
        _add_meta(doc, "강좌 코드", course.course_code)
    if course.professor:
        _add_meta(doc, "교수", course.professor)
    if course.semester:
        _add_meta(doc, "학기", course.semester)
    _add_meta(doc, "출처", source_label or paths.FALLBACK_SOURCE)
    _add_meta(doc, "생성일", datetime.now().strftime("%Y-%m-%d %H:%M"))


def _write_parsed_body(doc: Document, material: Material, pc: Optional[ParsedContent]) -> None:
    if pc is None:
        doc.add_paragraph("(파싱 대기 — 다음 동기화에서 본문이 채워집니다.)")
        return

    blocks = _load_blocks(pc)
    if not blocks:
        if material.parse_status == "skipped":
            doc.add_paragraph(f"(파싱 미지원 형식: {material.file_type or '?'})")
        elif material.parse_status == "failed":
            doc.add_paragraph(f"(파싱 실패: {material.parse_error or '알 수 없는 오류'})")
        else:
            doc.add_paragraph("(추출된 본문이 없습니다.)")
        return

    if pc.used_ocr:
        meta = doc.add_paragraph()
        run = meta.add_run("(OCR 인식 결과 — 정확도 점검 필요)")
        run.italic = True

    last_page = None
    for b in blocks:
        text = (b.get("text") or "").strip()
        if not text:
            continue
        page = b.get("page")
        if page != last_page:
            head = doc.add_paragraph()
            head_run = head.add_run(f"[p.{page}]")
            head_run.italic = True
            head_run.bold = True
            last_page = page
        doc.add_paragraph(text)


# ---------------------------------------------------------------------------
# Per-artifact generators

def generate_material_docx(
    db: Session,
    *,
    course: Course,
    material: Material,
) -> Optional[Path]:
    """Write one DOCX next to the original file in <source>/정리/<stem>.docx.

    Returns the saved path, or None when there's no parsed body yet (we still
    emit a placeholder file so the user can find it in the source folder)."""
    out_dir = paths.summary_dir(course.course_name, material.source_label)
    out_dir.mkdir(parents=True, exist_ok=True)

    filename = Path(material.file_path).name if material.file_path else None
    stem = Path(filename).stem if filename else (material.title or "untitled")
    out_path = out_dir / f"{_safe_filename(stem)}_정리.docx"

    doc = Document()
    _set_korean_default_font(doc)

    title_text = material.title or filename or "(제목 없음)"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _course_header(doc, course, material.source_label)
    if filename:
        _add_meta(doc, "파일", filename)
    type_tag = f" [{material.file_type.upper()}]" if material.file_type else ""
    size_tag = f" ({_human_size(material.size_bytes)})" if material.size_bytes else ""
    if type_tag or size_tag:
        _add_meta(doc, "형식", (type_tag + size_tag).strip())
    if material.week:
        _add_meta(doc, "주차", f"{material.week}주차")
        topic = _course_week_topic(db, course.id, material.week)
        if topic and (topic.get("topic") or "").strip():
            _add_meta(doc, "주제", topic["topic"].strip())

    doc.add_paragraph()

    _add_section_heading(doc, "파싱 본문")
    pc = _parsed_for_material(db, material.id)
    _write_parsed_body(doc, material, pc)

    doc.save(str(out_path))
    log.info("material DOCX written: %s", out_path)
    return out_path


def generate_notice_docx(
    db: Session,
    *,
    course: Course,
    notice: Notice,
) -> Optional[Path]:
    """Write one DOCX per board article into <source>/정리/공지_<title>.docx."""
    out_dir = paths.summary_dir(course.course_name, notice.source_label)
    out_dir.mkdir(parents=True, exist_ok=True)

    title_text = (notice.title or "공지").strip()
    out_path = out_dir / f"공지_{_safe_filename(title_text, fallback=f'bwid_{notice.bwid}')}.docx"

    doc = Document()
    _set_korean_default_font(doc)

    heading = doc.add_heading(title_text, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _course_header(doc, course, notice.source_label)
    if notice.author:
        _add_meta(doc, "작성자", notice.author)
    if notice.posted_at:
        _add_meta(doc, "게시일", notice.posted_at.strftime("%Y-%m-%d %H:%M"))
    if notice.url:
        _add_meta(doc, "원문 URL", notice.url)

    doc.add_paragraph()
    _add_section_heading(doc, "본문")

    body_text = _html_to_text(notice.body_html)
    if body_text:
        for para in body_text.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
    else:
        doc.add_paragraph("(본문이 비어 있거나 추출되지 않았습니다.)")

    doc.save(str(out_path))
    log.info("notice DOCX written: %s", out_path)
    return out_path


def generate_assignment_docx(
    db: Session,
    *,
    course: Course,
    assignment: Assignment,
) -> Optional[Path]:
    """Write one DOCX per assignment into <source>/정리/과제_<title>.docx."""
    out_dir = paths.summary_dir(course.course_name, assignment.source_label or assignment.title)
    out_dir.mkdir(parents=True, exist_ok=True)

    title_text = (assignment.title or "과제").strip()
    out_path = out_dir / f"과제_{_safe_filename(title_text, fallback=f'cmid_{assignment.cmid}')}.docx"

    doc = Document()
    _set_korean_default_font(doc)

    heading = doc.add_heading(title_text, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _course_header(doc, course, assignment.source_label or assignment.title)
    if assignment.due_at:
        _add_meta(doc, "마감", assignment.due_at.strftime("%Y-%m-%d %H:%M"))
    _add_meta(doc, "제출 상태", "제출 완료" if assignment.submitted else "미제출")
    if assignment.url:
        _add_meta(doc, "원문 URL", assignment.url)

    doc.add_paragraph()
    _add_section_heading(doc, "과제 설명")

    desc_text = _html_to_text(assignment.description_html)
    if desc_text:
        for para in desc_text.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
    else:
        doc.add_paragraph("(설명이 비어 있거나 추출되지 않았습니다.)")

    # ---- 첨부 자료 (이 과제 cmid로 연결된 Material 본문) ----
    attached = [m for m in _materials_for_course(db, course.id) if m.assignment_id == assignment.id]
    if attached:
        doc.add_paragraph()
        _add_section_heading(doc, "첨부 파일 본문")
        for m in attached:
            doc.add_heading(m.title or (Path(m.file_path).name if m.file_path else ""), level=2)
            pc = _parsed_for_material(db, m.id)
            _write_parsed_body(doc, m, pc)
            doc.add_paragraph()

    doc.save(str(out_path))
    log.info("assignment DOCX written: %s", out_path)
    return out_path


# ---------------------------------------------------------------------------
# Course-level aggregator

def generate_course_summaries(db: Session, *, course: Course) -> List[Path]:
    """Generate one DOCX per file/notice/assignment for this course.

    Returns the list of written paths."""
    written: List[Path] = []

    # 1) one DOCX per material (PDF / PPT / DOCX downloads)
    for m in _materials_for_course(db, course.id):
        try:
            p = generate_material_docx(db, course=course, material=m)
            if p is not None:
                written.append(p)
        except Exception:
            log.exception("material DOCX failed: course=%s material=%s", course.id, m.id)

    # 2) one DOCX per notice (board article body)
    for n in _notices_for_course(db, course.id):
        try:
            p = generate_notice_docx(db, course=course, notice=n)
            if p is not None:
                written.append(p)
        except Exception:
            log.exception("notice DOCX failed: course=%s notice=%s", course.id, n.id)

    # 3) one DOCX per assignment (description + attached parsed bodies)
    for a in _assignments_for_course(db, course.id):
        try:
            p = generate_assignment_docx(db, course=course, assignment=a)
            if p is not None:
                written.append(p)
        except Exception:
            log.exception("assignment DOCX failed: course=%s assignment=%s", course.id, a.id)

    return written


def generate_all(db: Session) -> dict:
    """Re-render summaries for every course in the DB. Returns counts."""
    courses = db.query(Course).all()
    total = 0
    for c in courses:
        total += len(generate_course_summaries(db, course=c))
    return {"courses": len(courses), "docx_written": total}
